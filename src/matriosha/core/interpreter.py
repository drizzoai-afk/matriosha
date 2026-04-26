"""Semantic content interpreter for agent-ready memory recall.

Single entrypoint: ``decode_semantic_content``.
It accepts payload bytes/base64 + metadata hints and returns a bounded,
structured semantic JSON object for immediate agent consumption.
"""

from __future__ import annotations

import base64
import binascii
import csv
import io
import json
import mimetypes
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from PIL import Image
from pypdf import PdfReader
import pytesseract
from docx import Document
from openpyxl import load_workbook

from matriosha.core.interpreter_plugins import REGISTRY


@dataclass(frozen=True)
class InterpreterBounds:
    """Deterministic extraction limits to prevent memory blowups."""

    max_input_bytes: int = 50 * 1024 * 1024
    max_preview_chars: int = 4096
    max_text_chars: int = 200_000
    max_tables: int = 16
    max_rows_per_table: int = 200
    max_cols_per_table: int = 64
    max_pdf_pages: int = 64
    max_doc_sections: int = 2000
    max_image_pixels: int = 25_000_000  # ~5000 x 5000


DEFAULT_BOUNDS = InterpreterBounds()


_TEXT_MIMES = {
    "text/plain",
    "text/markdown",
    "text/csv",
    "text/tab-separated-values",
    "application/json",
}


_SPREADSHEET_EXTS = {".xlsx"}
_TEXT_EXTS = {".txt", ".md", ".markdown", ".json", ".csv", ".tsv"}
_DOCUMENT_EXTS = {".docx"}
_LEGACY_DOCUMENT_EXTS = {".doc", ".odt"}
_LEGACY_SPREADSHEET_EXTS = {".xls"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tiff", ".tif"}


class _PdfDecoder:
    name = "builtin.pdf"

    def supports(self, mime_type: str, filename: str | None, metadata: dict[str, Any]) -> bool:
        ext = Path(filename).suffix.lower() if filename else ""
        return mime_type == "application/pdf" or ext == ".pdf"

    def decode(self, raw: bytes, metadata: dict[str, Any], bounds: InterpreterBounds) -> dict[str, Any]:
        semantic = _empty_semantic_patch(kind="pdf")
        _extract_pdf(raw, semantic, bounds)
        return semantic


class _ImageDecoder:
    name = "builtin.image"

    def supports(self, mime_type: str, filename: str | None, metadata: dict[str, Any]) -> bool:
        ext = Path(filename).suffix.lower() if filename else ""
        return mime_type.startswith("image/") or ext in _IMAGE_EXTS

    def decode(self, raw: bytes, metadata: dict[str, Any], bounds: InterpreterBounds) -> dict[str, Any]:
        semantic = _empty_semantic_patch(kind="image")
        _extract_image(raw, semantic, bounds)
        return semantic


class _TextDecoder:
    name = "builtin.text"

    def supports(self, mime_type: str, filename: str | None, metadata: dict[str, Any]) -> bool:
        ext = Path(filename).suffix.lower() if filename else ""
        return mime_type in _TEXT_MIMES or ext in _TEXT_EXTS

    def decode(self, raw: bytes, metadata: dict[str, Any], bounds: InterpreterBounds) -> dict[str, Any]:
        semantic = _empty_semantic_patch(kind="text")
        _extract_text(
            raw,
            semantic,
            bounds,
            mime_type=metadata.get("mime_type", "application/octet-stream"),
            filename=metadata.get("filename"),
        )
        return semantic


class _DocumentDecoder:
    name = "builtin.document"

    def supports(self, mime_type: str, filename: str | None, metadata: dict[str, Any]) -> bool:
        ext = Path(filename).suffix.lower() if filename else ""
        return ext in _DOCUMENT_EXTS or ext in _LEGACY_DOCUMENT_EXTS or mime_type in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }

    def decode(self, raw: bytes, metadata: dict[str, Any], bounds: InterpreterBounds) -> dict[str, Any]:
        semantic = _empty_semantic_patch(kind="document")
        _extract_document(raw, semantic, bounds, filename=metadata.get("filename"))
        return semantic


class _TableDecoder:
    name = "builtin.table"

    def supports(self, mime_type: str, filename: str | None, metadata: dict[str, Any]) -> bool:
        ext = Path(filename).suffix.lower() if filename else ""
        return ext in _SPREADSHEET_EXTS or ext in _LEGACY_SPREADSHEET_EXTS or mime_type in {
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
            "application/excel",
            "application/x-excel",
            "application/x-msexcel",
            # CSV/TSV are handled by the text decoder, which also extracts tables.
        }

    def decode(self, raw: bytes, metadata: dict[str, Any], bounds: InterpreterBounds) -> dict[str, Any]:
        semantic = _empty_semantic_patch(kind="table")
        _extract_table(
            raw,
            semantic,
            bounds,
            mime_type=metadata.get("mime_type", "application/octet-stream"),
            filename=metadata.get("filename"),
        )
        return semantic


class _BinaryFallbackDecoder:
    name = "builtin.binary_fallback"

    def supports(self, mime_type: str, filename: str | None, metadata: dict[str, Any]) -> bool:
        return True

    def decode(self, raw: bytes, metadata: dict[str, Any], bounds: InterpreterBounds) -> dict[str, Any]:
        semantic = _empty_semantic_patch(kind="binary")
        _extract_unknown(raw, semantic, bounds)
        return semantic


def _builtin_decoders() -> list[tuple[Any, str]]:
    return [
        (_PdfDecoder(), "builtin"),
        (_ImageDecoder(), "builtin"),
        (_TextDecoder(), "builtin"),
        (_DocumentDecoder(), "builtin"),
        (_TableDecoder(), "builtin"),
        (_BinaryFallbackDecoder(), "fallback"),
    ]


def _configure_registry_defaults() -> None:
    REGISTRY.set_default_factory(_builtin_decoders)
    REGISTRY.reset_default_decoders_for_tests()


def decode_semantic_content(
    payload: bytes | str,
    metadata: dict[str, Any] | None = None,
    *,
    bounds: InterpreterBounds = DEFAULT_BOUNDS,
) -> dict[str, Any]:
    """Decode payload into rich semantic JSON.

    Args:
        payload: Raw bytes or base64-encoded string/bytes.
        metadata: Optional hints, e.g. mime_type, filename, hints.
        bounds: Deterministic extraction bounds.
    """

    meta = dict(metadata or {})
    filename = _safe_filename(meta.get("filename"))
    mime_type = _resolve_mime_type(meta.get("mime_type"), filename)
    meta["filename"] = filename
    meta["mime_type"] = mime_type

    warnings: list[str] = []
    warnings.extend(REGISTRY.pull_warnings())

    raw = _resolve_payload_bytes(payload, warnings=warnings)
    if len(raw) > bounds.max_input_bytes:
        warnings.append(
            f"payload exceeded {bounds.max_input_bytes} bytes and was truncated for semantic extraction"
        )
        raw = raw[: bounds.max_input_bytes]

    semantic: dict[str, Any] = {
        "kind": "binary",
        "mime_type": mime_type,
        "filename": filename,
        "text": "",
        "tables": [],
        "metadata": {
            "input_bytes": len(raw),
            "extraction_bounds": {
                "max_preview_chars": bounds.max_preview_chars,
                "max_text_chars": bounds.max_text_chars,
                "max_tables": bounds.max_tables,
                "max_rows_per_table": bounds.max_rows_per_table,
                "max_cols_per_table": bounds.max_cols_per_table,
            },
        },
        "warnings": warnings,
        "preview": "",
    }

    matching = REGISTRY.get_matching_decoders(mime_type, filename, meta)
    warnings.extend(REGISTRY.pull_warnings())

    non_fallback_matches = [plugin for plugin in matching if plugin.name != "builtin.binary_fallback"]
    if len(non_fallback_matches) > 1:
        names = [plugin.name for plugin in non_fallback_matches]
        warnings.append(
            f"multiple decoder plugins matched; using '{names[0]}' (deterministic priority), skipped: {', '.join(names[1:])}"
        )

    selected = matching[0] if matching else _BinaryFallbackDecoder()

    try:
        patch = selected.decode(raw, dict(meta), bounds)
        REGISTRY.increment_usage(selected.name)
        _merge_semantic_patch(semantic, patch)
    except Exception as exc:  # noqa: BLE001
        semantic["warnings"].append(
            f"semantic extraction failed: {type(exc).__name__}: {exc}"
        )
        fallback_semantic = _empty_semantic_patch(kind="binary")
        _extract_unknown(raw, fallback_semantic, bounds)
        _merge_semantic_patch(semantic, fallback_semantic)

    text_value = str(semantic.get("text") or "")
    semantic["text"] = _clip_text(text_value, bounds.max_text_chars, semantic["warnings"])
    semantic["preview"] = _build_preview(semantic["text"], bounds.max_preview_chars)
    return semantic


def _empty_semantic_patch(*, kind: str) -> dict[str, Any]:
    return {
        "kind": kind,
        "text": "",
        "tables": [],
        "metadata": {},
        "warnings": [],
    }


def _merge_semantic_patch(semantic: dict[str, Any], patch: dict[str, Any]) -> None:
    kind = patch.get("kind")
    if isinstance(kind, str) and kind:
        semantic["kind"] = kind

    text = patch.get("text")
    if text is not None:
        semantic["text"] = str(text)

    tables = patch.get("tables")
    if isinstance(tables, list):
        semantic["tables"] = tables

    metadata_patch = patch.get("metadata")
    if isinstance(metadata_patch, dict):
        semantic["metadata"].update(metadata_patch)

    warning_patch = patch.get("warnings")
    if isinstance(warning_patch, list):
        semantic["warnings"].extend(str(w) for w in warning_patch)


def _resolve_payload_bytes(payload: bytes | str, *, warnings: list[str]) -> bytes:
    if isinstance(payload, bytes):
        return payload

    if isinstance(payload, str):
        stripped = payload.strip()
        if not stripped:
            return b""
        try:
            return base64.b64decode(stripped, validate=True)
        except (ValueError, binascii.Error):
            warnings.append("payload string is not valid base64; treated as utf-8 text bytes")
            return stripped.encode("utf-8", errors="replace")

    raise TypeError("payload must be bytes or base64 string")


def _resolve_mime_type(raw_mime: Any, filename: str | None) -> str:
    if isinstance(raw_mime, str) and raw_mime.strip():
        return raw_mime.strip().lower()

    if filename:
        guessed, _ = mimetypes.guess_type(filename)
        if guessed:
            return guessed.lower()

    return "application/octet-stream"


def _safe_filename(value: Any) -> str | None:
    if not isinstance(value, str) or not value.strip():
        return None
    return Path(value).name


def _clip_text(text: str, limit: int, warnings: list[str]) -> str:
    if len(text) <= limit:
        return text
    warnings.append(f"text truncated to {limit} characters")
    return text[:limit]


def _build_preview(text: str, max_chars: int) -> str:
    normalized = " ".join(text.split())
    if len(normalized) <= max_chars:
        return normalized
    return normalized[:max_chars]


def _extract_pdf(raw: bytes, semantic: dict[str, Any], bounds: InterpreterBounds) -> None:
    reader = PdfReader(io.BytesIO(raw))
    total_pages = len(reader.pages)
    page_limit = min(total_pages, bounds.max_pdf_pages)
    pages_meta: list[dict[str, Any]] = []
    text_chunks: list[str] = []

    for idx in range(page_limit):
        page = reader.pages[idx]
        text = (page.extract_text() or "").strip()
        text_chunks.append(text)
        table_candidates = _table_candidate_count(text)
        pages_meta.append(
            {
                "page": idx + 1,
                "chars": len(text),
                "table_candidate_lines": table_candidates,
            }
        )

    if total_pages > page_limit:
        semantic["warnings"].append(f"pdf pages truncated to {page_limit}/{total_pages}")

    semantic["metadata"]["page_count"] = total_pages
    semantic["metadata"]["pages"] = pages_meta
    semantic["metadata"]["table_candidates"] = sum(page["table_candidate_lines"] for page in pages_meta)
    semantic["text"] = "\n\n".join(chunk for chunk in text_chunks if chunk)


def _extract_image(raw: bytes, semantic: dict[str, Any], bounds: InterpreterBounds) -> None:
    with Image.open(io.BytesIO(raw)) as img:
        width, height = img.size
        semantic["metadata"]["width"] = width
        semantic["metadata"]["height"] = height
        semantic["metadata"]["format"] = img.format
        semantic["metadata"]["mode"] = img.mode

        pixels = width * height
        working = img
        if pixels > bounds.max_image_pixels:
            factor = (bounds.max_image_pixels / float(pixels)) ** 0.5
            new_size = (
                max(1, int(width * factor)),
                max(1, int(height * factor)),
            )
            working = img.copy()
            working.thumbnail(new_size)
            semantic["warnings"].append(
                f"image downscaled from {width}x{height} to {working.size[0]}x{working.size[1]} for OCR bounds"
            )

        try:
            ocr_text = pytesseract.image_to_string(working)
        except Exception as exc:  # noqa: BLE001
            semantic["warnings"].append(f"ocr unavailable: {type(exc).__name__}: {exc}")
            ocr_text = ""

        semantic["text"] = (ocr_text or "").strip()


def _extract_text(
    raw: bytes,
    semantic: dict[str, Any],
    bounds: InterpreterBounds,
    *,
    mime_type: str,
    filename: str | None,
) -> None:
    decoded = raw.decode("utf-8", errors="replace")
    normalized = decoded.replace("\r\n", "\n").replace("\r", "\n")

    if mime_type == "application/json" or (filename and filename.lower().endswith(".json")):
        try:
            obj = json.loads(normalized)
            normalized = json.dumps(obj, indent=2, ensure_ascii=False, sort_keys=True)
            semantic["metadata"]["json_valid"] = True
        except Exception:  # noqa: BLE001
            semantic["metadata"]["json_valid"] = False

    semantic["metadata"]["line_count"] = normalized.count("\n") + 1 if normalized else 0
    semantic["text"] = normalized

    if filename and filename.lower().endswith((".csv", ".tsv")):
        _extract_table(raw, semantic, bounds, mime_type=mime_type, filename=filename)


def _extract_document(raw: bytes, semantic: dict[str, Any], bounds: InterpreterBounds, *, filename: str | None) -> None:
    ext = Path(filename).suffix.lower() if filename else ""
    if ext in _LEGACY_DOCUMENT_EXTS:
        semantic["kind"] = "binary"
        semantic["warnings"].append(f"{ext} is not rich-decoded by the built-in document decoder")
        _extract_unknown(raw, semantic, bounds)
        return

    doc = Document(io.BytesIO(raw))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    if len(paragraphs) > bounds.max_doc_sections:
        semantic["warnings"].append(
            f"document paragraphs truncated to {bounds.max_doc_sections}/{len(paragraphs)}"
        )
        paragraphs = paragraphs[: bounds.max_doc_sections]

    semantic["metadata"]["section_count"] = len(paragraphs)
    semantic["metadata"]["doc_table_count"] = len(doc.tables)
    semantic["text"] = "\n".join(paragraphs)

    tables = []
    for doc_table in doc.tables[: bounds.max_tables]:
        rows: list[list[str]] = []
        for row in doc_table.rows[: bounds.max_rows_per_table]:
            cells = [str(cell.text).strip() for cell in row.cells[: bounds.max_cols_per_table]]
            rows.append(cells)
        tables.append(
            {
                "name": "doc_table",
                "row_count": len(rows),
                "column_count": max((len(r) for r in rows), default=0),
                "rows": rows,
            }
        )
    semantic["tables"] = tables


def _extract_table(
    raw: bytes,
    semantic: dict[str, Any],
    bounds: InterpreterBounds,
    *,
    mime_type: str,
    filename: str | None,
) -> None:
    ext = Path(filename).suffix.lower() if filename else ""

    if ext in _LEGACY_SPREADSHEET_EXTS:
        semantic["kind"] = "binary"
        semantic["warnings"].append(f"{ext} is not rich-decoded by the built-in table decoder")
        _extract_unknown(raw, semantic, bounds)
        return

    if ext in {".csv", ".tsv"} or mime_type in {"text/csv", "text/tab-separated-values"}:
        delimiter = "\t" if ext == ".tsv" or mime_type == "text/tab-separated-values" else ","
        _extract_delimited_table(raw, semantic, bounds, delimiter=delimiter)
        return

    workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    tables: list[dict[str, Any]] = []
    text_parts: list[str] = []

    for sheet_idx, sheet_name in enumerate(workbook.sheetnames):
        if sheet_idx >= bounds.max_tables:
            semantic["warnings"].append(f"worksheet count truncated to {bounds.max_tables}")
            break

        sheet = workbook[sheet_name]
        rows: list[list[str]] = []
        max_cols_seen = 0
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
            if row_idx >= bounds.max_rows_per_table:
                semantic["warnings"].append(
                    f"worksheet '{sheet_name}' rows truncated to {bounds.max_rows_per_table}"
                )
                break
            trimmed = ["" if cell is None else str(cell) for cell in row[: bounds.max_cols_per_table]]
            max_cols_seen = max(max_cols_seen, len(trimmed))
            rows.append(trimmed)

        table = {
            "name": sheet_name,
            "row_count": len(rows),
            "column_count": max_cols_seen,
            "rows": rows,
        }
        tables.append(table)

        header = rows[0] if rows else []
        preview_rows = rows[1: min(len(rows), 6)]
        text_parts.append(f"Sheet: {sheet_name}")
        if header:
            text_parts.append(" | ".join(header))
        for prow in preview_rows:
            text_parts.append(" | ".join(prow))

    semantic["tables"] = tables
    semantic["metadata"]["sheet_count"] = len(workbook.sheetnames)
    semantic["text"] = "\n".join(text_parts)


def _extract_delimited_table(
    raw: bytes,
    semantic: dict[str, Any],
    bounds: InterpreterBounds,
    *,
    delimiter: str,
) -> None:
    decoded = raw.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(decoded, newline=""), delimiter=delimiter)

    rows: list[list[str]] = []
    max_cols_seen = 0
    for idx, row in enumerate(reader):
        if idx >= bounds.max_rows_per_table:
            semantic["warnings"].append(f"table rows truncated to {bounds.max_rows_per_table}")
            break
        trimmed = [str(cell) for cell in row[: bounds.max_cols_per_table]]
        max_cols_seen = max(max_cols_seen, len(trimmed))
        rows.append(trimmed)

    semantic["tables"] = [
        {
            "name": "table",
            "row_count": len(rows),
            "column_count": max_cols_seen,
            "rows": rows,
        }
    ]

    header = rows[0] if rows else []
    preview_rows = rows[1: min(len(rows), 6)]
    text_lines = []
    if header:
        text_lines.append(" | ".join(header))
    for prow in preview_rows:
        text_lines.append(" | ".join(prow))
    semantic["text"] = "\n".join(text_lines)


def _extract_unknown(raw: bytes, semantic: dict[str, Any], bounds: InterpreterBounds) -> None:
    hex_preview = raw[:32].hex()

    semantic["metadata"]["binary_preview_hex"] = hex_preview
    semantic["metadata"]["is_probably_text"] = _is_probably_text(raw)
    semantic["warnings"].append("unknown binary payload; returning bounded hex preview only")
    semantic["text"] = ""


def _is_probably_text(raw: bytes) -> bool:
    if not raw:
        return True
    sample = raw[:1024]
    text_like = sum(1 for b in sample if 9 <= b <= 13 or 32 <= b <= 126)
    return (text_like / len(sample)) > 0.85


def _table_candidate_count(text: str) -> int:
    count = 0
    for line in text.splitlines():
        if line.count("|") >= 2 or line.count("\t") >= 2:
            count += 1
    return count


_configure_registry_defaults()
