from __future__ import annotations
import base64

import io
import json

from PIL import Image
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from matriosha.core.interpreter import InterpreterBounds, decode_semantic_content


def test_decode_text_json_and_csv_contract_fields() -> None:
    payload = json.dumps({"hello": "world", "n": 1}).encode("utf-8")
    semantic = decode_semantic_content(payload, {"mime_type": "application/json", "filename": "sample.json"})

    assert semantic["kind"] == "text"
    assert semantic["mime_type"] == "application/json"
    assert semantic["filename"] == "sample.json"
    assert isinstance(semantic["text"], str)
    assert isinstance(semantic["tables"], list)
    assert isinstance(semantic["metadata"], dict)
    assert isinstance(semantic["warnings"], list)
    assert isinstance(semantic["preview"], str)
    assert semantic["metadata"]["json_valid"] is True

    csv_semantic = decode_semantic_content(
        b"name,score\nalice,10\nbob,9\n",
        {"mime_type": "text/csv", "filename": "scores.csv"},
    )
    assert csv_semantic["kind"] == "text"
    assert len(csv_semantic["tables"]) == 1
    assert csv_semantic["tables"][0]["row_count"] == 3


def test_decode_pdf_docx_xlsx_best_effort() -> None:
    pdf_writer = PdfWriter()
    pdf_writer.add_blank_page(width=200, height=200)
    pdf_buffer = io.BytesIO()
    pdf_writer.write(pdf_buffer)
    pdf_semantic = decode_semantic_content(pdf_buffer.getvalue(), {"filename": "f.pdf"})
    assert pdf_semantic["kind"] == "pdf"
    assert pdf_semantic["metadata"]["page_count"] == 1

    doc = Document()
    doc.add_paragraph("alpha section")
    doc.add_paragraph("beta section")
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_semantic = decode_semantic_content(
        doc_buffer.getvalue(),
        {"filename": "memo.docx"},
    )
    assert doc_semantic["kind"] == "document"
    assert "alpha section" in doc_semantic["text"]

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["id", "value"])
    ws.append([1, "A"])
    xlsx_buffer = io.BytesIO()
    wb.save(xlsx_buffer)
    xlsx_semantic = decode_semantic_content(
        xlsx_buffer.getvalue(),
        {"filename": "table.xlsx"},
    )
    assert xlsx_semantic["kind"] == "table"
    assert xlsx_semantic["metadata"]["sheet_count"] == 1
    assert xlsx_semantic["tables"][0]["column_count"] >= 2


def test_decode_image_ocr_and_unknown_fallback(monkeypatch) -> None:
    monkeypatch.setattr("matriosha.core.interpreter.pytesseract.image_to_string", lambda _img: "OCR text here")

    img = Image.new("RGB", (40, 30), color=(255, 255, 255))
    img_buf = io.BytesIO()
    img.save(img_buf, format="PNG")
    image_semantic = decode_semantic_content(img_buf.getvalue(), {"filename": "image.png"})

    assert image_semantic["kind"] == "image"
    assert image_semantic["metadata"]["width"] == 40
    assert image_semantic["metadata"]["height"] == 30
    assert "OCR text here" in image_semantic["text"]

    unknown = decode_semantic_content(b"\x00\x01\x02\xff\xf0", {"filename": "blob.bin"})
    assert unknown["kind"] == "binary"
    assert unknown["metadata"]["input_bytes"] == 5
    assert unknown["warnings"]


def test_interpreter_bounds_are_deterministic() -> None:
    semantic = decode_semantic_content(
        ("A" * 1000).encode("utf-8"),
        {"mime_type": "text/plain", "filename": "big.txt"},
        bounds=InterpreterBounds(max_text_chars=120, max_preview_chars=20),
    )

    assert len(semantic["text"]) == 120
    assert len(semantic["preview"]) <= 20
    assert any("truncated" in warning for warning in semantic["warnings"])



def test_interpreter_accepts_base64_payload_and_infers_mime_from_filename() -> None:
    payload = base64.b64encode(b"title,owner\nlaunch,daniele\n").decode("ascii")

    semantic = decode_semantic_content(payload, {"filename": "plan.csv"})

    assert semantic["mime_type"] == "text/csv"
    assert semantic["filename"] == "plan.csv"
    assert semantic["tables"][0]["row_count"] == 2
    assert "launch" in semantic["text"]


def test_interpreter_truncates_large_input_before_extraction() -> None:
    semantic = decode_semantic_content(
        b"A" * 100,
        {"mime_type": "text/plain", "filename": "large.txt"},
        bounds=InterpreterBounds(max_input_bytes=10, max_text_chars=100, max_preview_chars=100),
    )

    assert semantic["metadata"]["input_bytes"] == 10
    assert semantic["text"] == "A" * 10
    assert any("payload exceeded" in warning for warning in semantic["warnings"])


def test_interpreter_csv_row_and_column_bounds_warn() -> None:
    raw = b"a,b,c,d\n1,2,3,4\n5,6,7,8\n"

    semantic = decode_semantic_content(
        raw,
        {"filename": "bounded.csv"},
        bounds=InterpreterBounds(max_rows_per_table=2, max_cols_per_table=2),
    )

    table = semantic["tables"][0]
    assert table["row_count"] == 2
    assert table["column_count"] == 2
    assert table["rows"] == [["a", "b"], ["1", "2"]]
    assert any("table rows truncated" in warning for warning in semantic["warnings"])


def test_interpreter_corrupt_known_file_falls_back_without_crashing() -> None:
    semantic = decode_semantic_content(b"not a real pdf", {"filename": "broken.pdf"})

    assert semantic["kind"] == "binary"
    assert semantic["metadata"]["is_probably_text"] is True
    assert any("semantic extraction failed" in warning for warning in semantic["warnings"])


def test_interpreter_ocr_unavailable_warns_without_crashing(monkeypatch) -> None:
    def fail_ocr(_img):
        raise RuntimeError("tesseract missing")

    monkeypatch.setattr("matriosha.core.interpreter.pytesseract.image_to_string", fail_ocr)

    img = Image.new("RGB", (20, 20), color=(255, 255, 255))
    img_buf = io.BytesIO()
    img.save(img_buf, format="PNG")

    semantic = decode_semantic_content(img_buf.getvalue(), {"filename": "ocr.png"})

    assert semantic["kind"] == "image"
    assert semantic["text"] == ""
    assert any("ocr unavailable" in warning for warning in semantic["warnings"])


def test_interpreter_zip_uses_bounded_binary_fallback() -> None:
    semantic = decode_semantic_content(
        b"PK\x03\x04" + b"\x00" * 5000,
        {"filename": "archive.zip"},
        bounds=InterpreterBounds(max_preview_chars=64),
    )

    assert semantic["kind"] == "binary"
    assert semantic["mime_type"] == "application/zip"
    assert len(semantic["preview"]) <= 64
    assert "binary_preview_hex" in semantic["metadata"]
    assert any("unknown binary payload" in warning for warning in semantic["warnings"])



def test_interpreter_file_type_matrix_and_fallback_truthfulness() -> None:
    # Plain text
    txt = decode_semantic_content(b"alpha beta", {"filename": "note.txt"})
    assert txt["kind"] == "text"
    assert txt["text"] == "alpha beta"
    assert not any("unknown binary payload" in w for w in txt["warnings"])

    # Markdown
    md = decode_semantic_content(b"# Title\n\nLaunch notes", {"filename": "notes.md"})
    assert md["kind"] == "text"
    assert "Launch notes" in md["text"]

    # CSV rich table extraction
    csv_semantic = decode_semantic_content(b"name,score\nalice,10\n", {"filename": "scores.csv"})
    assert csv_semantic["kind"] == "text"
    assert csv_semantic["tables"]
    assert csv_semantic["tables"][0]["row_count"] == 2

    # DOCX rich document extraction
    doc = Document()
    doc.add_paragraph("Roadmap section")
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    docx_semantic = decode_semantic_content(doc_buffer.getvalue(), {"filename": "roadmap.docx"})
    assert docx_semantic["kind"] == "document"
    assert "Roadmap section" in docx_semantic["text"]
    assert docx_semantic["metadata"]["section_count"] == 1

    # XLSX rich spreadsheet extraction
    wb = Workbook()
    ws = wb.active
    ws.append(["feature", "status"])
    ws.append(["memory", "green"])
    xlsx_buffer = io.BytesIO()
    wb.save(xlsx_buffer)
    xlsx_semantic = decode_semantic_content(xlsx_buffer.getvalue(), {"filename": "matrix.xlsx"})
    assert xlsx_semantic["kind"] == "table"
    assert xlsx_semantic["tables"]
    assert xlsx_semantic["metadata"]["sheet_count"] == 1

    # Unsupported binary fallback must be explicit and bounded.
    binary = decode_semantic_content(b"\x00\x01\x02\x03", {"filename": "blob.bin"})
    assert binary["kind"] == "binary"
    assert "binary_preview_hex" in binary["metadata"]
    assert any("unknown binary payload" in w for w in binary["warnings"])


def test_interpreter_unsupported_office_formats_do_not_claim_rich_extraction() -> None:
    legacy_doc = decode_semantic_content(b"legacy doc bytes", {"filename": "legacy.doc"})
    assert legacy_doc["kind"] == "binary"
    assert legacy_doc["mime_type"] == "application/msword"
    assert legacy_doc["tables"] == []
    assert "binary_preview_hex" in legacy_doc["metadata"]
    assert any("not rich-decoded" in w for w in legacy_doc["warnings"])
    assert any("unknown binary payload" in w for w in legacy_doc["warnings"])

    legacy_xls = decode_semantic_content(b"legacy xls bytes", {"filename": "legacy.xls"})
    assert legacy_xls["kind"] == "binary"
    assert legacy_xls["mime_type"] in {
        "application/vnd.ms-excel",
        "application/excel",
        "application/x-excel",
        "application/x-msexcel",
    }
    assert legacy_xls["tables"] == []
    assert "binary_preview_hex" in legacy_xls["metadata"]
    assert any("not rich-decoded" in w for w in legacy_xls["warnings"])
    assert any("unknown binary payload" in w for w in legacy_xls["warnings"])


def test_interpreter_archive_fallback_is_non_recursive_and_truthful() -> None:
    import zipfile

    archive_buffer = io.BytesIO()
    with zipfile.ZipFile(archive_buffer, "w") as archive:
        archive.writestr("inside.txt", "secret nested text that should not be extracted")

    semantic = decode_semantic_content(archive_buffer.getvalue(), {"filename": "archive.zip"})

    assert semantic["kind"] == "binary"
    assert semantic["mime_type"] == "application/zip"
    assert semantic["tables"] == []
    assert "binary_preview_hex" in semantic["metadata"]
    assert "secret nested text" not in semantic["text"]
    assert "inside.txt" not in semantic["text"]
    assert any("unknown binary payload" in w for w in semantic["warnings"])


def test_interpreter_broken_known_file_fallback_does_not_keep_rich_kind() -> None:
    semantic = decode_semantic_content(b"not really a docx", {"filename": "broken.docx"})

    assert semantic["kind"] == "binary"
    assert semantic["tables"] == []
    assert "binary_preview_hex" in semantic["metadata"]
    assert any("semantic extraction failed" in w for w in semantic["warnings"])
    assert any("unknown binary payload" in w for w in semantic["warnings"])
